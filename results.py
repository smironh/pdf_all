from pdfplumber import PDF
from pathlib import Path
from gtts import gTTS

from docx import Document
from docx.shared import Inches

import os

def pdf_to_txt(file_name):
	print('[+] Файл принят')
	print('[+] Конвертируем...')

	with PDF(open(file=file_name, mode='rb')) as pdf:
		pages = [page.extract_text() for page in pdf.pages]
			
	text = ''.join(pages)

	file_name =  Path(file_name).stem

	print(f'[+] Сохраянем результаты в {file_name}.txt...')
	with open(f'result/{file_name}.txt', 'w', encoding='utf-8') as file:
		file.write(text)

	print('[+] Готово!')

def pdf_to_mp3(file_path='test.pdf', language='en'):
	
	print('[+] Файл принят')
	print('[+] Конвертируем...')
		
	with PDF(open(file=file_path, mode='rb')) as pdf:
		pages = [page.extract_text() for page in pdf.pages]
			
	text = ''.join(pages)
	text = text.replace('\n', '')
		
	my_audio = gTTS(text=text, lang=language, slow=False)
	file_name = Path(file_path).stem
	my_audio.save(f'result/{file_name}.mp3')

	print('[+] Готово')

def pdf_to_html(file_path):
	print('[+] Файл принят')
	print('[+] Конвертируем...')

	with PDF(open(file=file_path, mode='rb')) as pdf:
		pages = [page.extract_text() for page in pdf.pages]

	text = ''.join(pages)
	a = text.split('\n')
	b = 0

	file_name =  Path(file_path).stem
	with open(f'result/{file_name}.html', 'w', encoding='utf-8') as file:
		file.write('''
<style>
	body{
	font-family: Arial;
	}
</style>
''')
	file = open(f'result/{file_name}.html', 'w', encoding='utf-8')
	for i in range(len(a)):
		txt = a[b]
		txt = f'<p>{txt}</p>'
		file.write(txt)
		b+=1
	file.close()

	print('[+] Готово!')

def pdf_to_docx(file_path):
	print('[+] Файл принят')
	print('[+] Конвертируем...')

	with PDF(open(file=file_path, mode='rb')) as pdf:
		pages = [page.extract_text() for page in pdf.pages]


	file_name =  Path(file_path).stem	

	text = ''.join(pages)
	a = text.split('\n')
	b = 0

	document = Document()
	document.add_heading(file_name, 0)

	for i in range(len(a)):
		txt = a[b]
		txt = f'{txt}'
		document.add_paragraph(txt)
		b+=1

	document.save(f'result/{file_name}.docx')
	print('[+] Готово!')
